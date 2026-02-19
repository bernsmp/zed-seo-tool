"""
SOW utilities — transcript extraction, DOCX generation, PM brief builder.
Matches TrackableMed's existing SOW format based on pattern analysis of 9 SOWs.
"""

import io
import json
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from utils.llm import _chat, _parse_json

# ── Constants ────────────────────────────────────────────────────

SERVICE_TYPES = {
    "seo_web": "SEO / Website Management",
    "digital_ads": "Digital Advertising",
    "practice_consulting": "Practice Consulting (PGA)",
}

REPS = ["Erik Trattler", "Zed Williamson"]

ACCEPTANCE_BODY = (
    "The client named below verifies that the terms of this Statement of Work are "
    "acceptable. The parties hereto are each acting with proper authority by their "
    "respective companies."
)

CANCELLATION_CLAUSE = (
    "Client(s) can cancel services with a 30-day written notice. Email notifications "
    "must be received and confirmed by TrackableMed thirty days preceding the desired "
    "cancellation date. Client(s) will be responsible for all associated monthly service "
    "fees leading up to the cancellation end date and/or until all essential actions have "
    "been completed by all parties during the exit process."
)

TERMINATION_CLAUSE = (
    "Termination shall not, under any circumstances, relieve Client of its obligation "
    "to pay any sums owed to Company under the terms of the agreement."
)


# ── LLM: Extract fields from transcript ─────────────────────────

def extract_from_transcript(transcript: str) -> dict:
    """
    Extract key SOW fields from a sales call transcript.
    Returns structured dict with client info, service type, signals.
    """
    prompt = f"""You are analyzing a sales call transcript between a TrackableMed sales rep and a prospective medical practice client.

Extract the following information from this transcript. Be specific and accurate — do not invent details not present in the transcript.

TRANSCRIPT:
{transcript[:8000]}

Return a JSON object:
{{
  "client_name": "Practice or business name (string, or null if unclear)",
  "client_contact": "Client's first and last name (string, or null if unclear)",
  "practice_specialty": "Medical specialty, e.g. ENT, Urology, Pain Management, Orthopedics, General Practice (string)",
  "service_type": "seo_web | digital_ads | practice_consulting — best match based on what was discussed",
  "service_type_confidence": 0-100,
  "service_type_reasoning": "One sentence explaining why you chose this service type",
  "services_mentioned": ["list of specific services, platforms, or deliverables mentioned"],
  "client_goals": "Brief summary of what the client wants to achieve (1-2 sentences, or null)",
  "budget_signals": "Any dollar amounts or budget ranges mentioned as a string, or null",
  "timeline_signals": "Any timeline or start date mentioned as a string, or null",
  "platforms": ["Google Ads", "META", "Connected TV", etc. — only what was explicitly mentioned],
  "notes": "Any other important context for the SOW (string, or null)"
}}"""

    raw = _chat(
        [{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        task="sow_extract",
    )
    return _parse_json(raw)


# ── LLM: Generate SOW narrative content ─────────────────────────

def generate_sow_content(fields: dict) -> dict:
    """
    Generate summary paragraphs, scope items, job name, and optional timeline
    from confirmed SOW fields.
    """
    service_type = fields.get("service_type", "seo_web")

    service_guidance = {
        "seo_web": (
            "Generate content for an SEO / Website Management engagement.\n\n"
            "scope_items should use numbered sections (matching MD Diagnostics SOW pattern) covering:\n"
            "- Keyword Research & Analysis (include specific target keyword count)\n"
            "- On-Page Optimization (list specific page types)\n"
            "- Technical SEO Audit & Implementation\n"
            "- Content Creation (blog posts per month if relevant)\n"
            "- Local SEO (if applicable based on practice type)\n"
            "- Link Building & Authority Building\n"
            "- Monthly Reporting & Analytics Review\n"
            "- Any website management or additional items mentioned\n\n"
            "timeline_items should be an empty array — SEO is ongoing from day one."
        ),
        "digital_ads": (
            "Generate content for a Digital Advertising engagement.\n\n"
            "scope_items should use two phases:\n"
            "Phase 1 — Asset Build: landing page creation, ad creative, pixel/tracking setup, audience targeting\n"
            "Phase 2 — Campaign Management: campaign launch, A/B testing, bid management, monthly reporting\n\n"
            "Also generate timeline_items with 4-5 activities and week ranges "
            "(e.g., 'Week 1-2', 'Week 3', 'Week 4+'). Include activities like:\n"
            "- Discovery & Strategy\n"
            "- Asset Build (landing pages, creative)\n"
            "- Pixel Setup & Testing\n"
            "- Campaign Launch\n"
            "- Ongoing Optimization"
        ),
        "practice_consulting": (
            "Generate content for a Practice Consulting (Physician Growth Accelerator) engagement.\n\n"
            "scope_items should cover:\n"
            "- Initial Practice Assessment / vitals audit\n"
            "- Monthly strategic consulting calls\n"
            "- Operational efficiency improvements\n"
            "- Patient acquisition and retention strategy\n"
            "- Any specific areas mentioned by the client\n\n"
            "Frame it around measurable outcomes (efficiency, patient volume, revenue growth).\n"
            "timeline_items: include 3-4 milestone activities."
        ),
    }

    services_str = ", ".join(fields.get("services_mentioned", [])) or "general services"
    platforms_str = ", ".join(fields.get("platforms", [])) or "relevant platforms"

    prompt = f"""You are a professional medical marketing copywriter for TrackableMed.
Write SOW content based on these confirmed engagement details:

CLIENT: {fields.get("client_name", "[Client Name]")}
SPECIALTY: {fields.get("practice_specialty", "Medical Practice")}
SERVICE TYPE: {SERVICE_TYPES.get(service_type, service_type)}
GOALS: {fields.get("client_goals", "Grow patient volume and practice revenue")}
SERVICES MENTIONED: {services_str}
PLATFORMS: {platforms_str}
NOTES: {fields.get("notes", "")}

{service_guidance.get(service_type, service_guidance["seo_web"])}

Return a JSON object:
{{
  "job_name": "Short service description, 3-6 words (e.g. 'SEO & Website Management 2026', 'Digital Advertising Campaign', 'Physician Growth Accelerator')",
  "summary": "2-3 paragraphs separated by double newlines. Paragraph 1: the client's current situation and opportunity. Paragraph 2: what TrackableMed will do and the strategic approach. Paragraph 3 (optional): expected outcomes. Reference the client's specialty and specific goals — no generic filler.",
  "scope_items": [
    {{
      "heading": "Section or phase heading (string)",
      "points": ["bullet point 1", "bullet point 2", "..."]
    }}
  ],
  "timeline_items": [
    {{"activity": "Activity name", "start": "Week X", "end": "Week Y"}}
  ]
}}

Scope should have 4-8 sections with 2-5 bullet points each. Write professionally — these go directly to medical practice owners."""

    raw = _chat(
        [{"role": "user", "content": prompt}],
        response_format={"type": "json_object"},
        task="sow_generate",
    )
    return _parse_json(raw)


# ── DOCX helpers ─────────────────────────────────────────────────

def _color_heading(heading, hex_r: int, hex_g: int, hex_b: int):
    """Apply RGB color to all runs in a heading paragraph."""
    for run in heading.runs:
        run.font.color.rgb = RGBColor(hex_r, hex_g, hex_b)


def _bold_cell(cell):
    """Bold all text in a table cell."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def _set_cell_width(cell, width_inches: float):
    """Set a table cell's preferred width."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_inches * 1440)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ── DOCX builder ─────────────────────────────────────────────────

def build_sow_docx(fields: dict, content: dict) -> bytes:
    """
    Build a DOCX matching TrackableMed's SOW format.
    fields: confirmed user inputs (client name, rep, pricing, etc.)
    content: AI-generated narrative (summary, scope, job name, timeline)
    Returns raw bytes for st.download_button.
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Title ─────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Statement of Work")
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x1B, 0x2B, 0x5E)
    doc.add_paragraph()

    # ── Header table ──────────────────────────────────────────────
    header_table = doc.add_table(rows=5, cols=2)
    header_table.style = "Table Grid"

    today = date.today().strftime("%B %d, %Y")
    job_name = content.get("job_name", fields.get("job_name", "Services 2026"))

    header_rows = [
        ("Date", fields.get("sow_date", today)),
        ("Client", fields.get("client_name", "")),
        ("Job Name", job_name),
        ("Requested by", fields.get("client_contact", "")),
        ("From", fields.get("rep_name", "Erik Trattler")),
    ]

    for i, (label, value) in enumerate(header_rows):
        row = header_table.rows[i]
        lc = row.cells[0]
        lc.text = label
        for para in lc.paragraphs:
            for run in para.runs:
                run.bold = True
        row.cells[1].text = str(value or "")

    doc.add_paragraph()

    # ── Summary ───────────────────────────────────────────────────
    h = doc.add_heading("Summary", level=1)
    _color_heading(h, 0x1B, 0x2B, 0x5E)

    summary_text = content.get("summary", "")
    for chunk in summary_text.split("\n\n"):
        if chunk.strip():
            doc.add_paragraph(chunk.strip())
    doc.add_paragraph()

    # ── Scope of Services ─────────────────────────────────────────
    h = doc.add_heading("Scope of Services", level=1)
    _color_heading(h, 0x1B, 0x2B, 0x5E)

    service_type = fields.get("service_type", "seo_web")
    scope_items = content.get("scope_items", [])

    for i, section in enumerate(scope_items):
        heading_text = section.get("heading", "")
        points = section.get("points", [])

        # Number sections for SEO (matches MD Diagnostics pattern)
        if service_type == "seo_web":
            label = f"2.{i + 1}  {heading_text}"
        else:
            label = heading_text

        h2 = doc.add_heading(label, level=2)
        _color_heading(h2, 0x2D, 0x39, 0x8E)

        for point in points:
            p = doc.add_paragraph(f"\u2022  {point}")
            p.paragraph_format.left_indent = Inches(0.25)

        doc.add_paragraph()

    # ── Schedule / Timeline (optional) ───────────────────────────
    timeline_items = content.get("timeline_items", [])
    if timeline_items:
        h = doc.add_heading("Schedule", level=1)
        _color_heading(h, 0x1B, 0x2B, 0x5E)

        tl = doc.add_table(rows=1, cols=3)
        tl.style = "Table Grid"
        hdr = tl.rows[0]
        hdr.cells[0].text = "Activity"
        hdr.cells[1].text = "Start"
        hdr.cells[2].text = "End"
        for cell in hdr.cells:
            _bold_cell(cell)

        for item in timeline_items:
            row = tl.add_row()
            row.cells[0].text = item.get("activity", "")
            row.cells[1].text = item.get("start", "")
            row.cells[2].text = item.get("end", "")

        doc.add_paragraph()

    # ── Pricing ───────────────────────────────────────────────────
    h = doc.add_heading("Pricing", level=1)
    _color_heading(h, 0x1B, 0x2B, 0x5E)
    doc.add_paragraph(
        "All costs listed below are based on the scope and assumptions "
        "included in this Statement of Work."
    )

    pricing_items = fields.get("pricing_items", [])
    if pricing_items:
        pt = doc.add_table(rows=1, cols=3)
        pt.style = "Table Grid"
        hdr = pt.rows[0]
        hdr.cells[0].text = "Item"
        hdr.cells[1].text = "Cost"
        hdr.cells[2].text = "Total"
        for cell in hdr.cells:
            _bold_cell(cell)

        for item in pricing_items:
            row = pt.add_row()
            row.cells[0].text = item.get("item", "")
            row.cells[1].text = item.get("cost", "")
            row.cells[2].text = item.get("total", "")

    payment_terms = fields.get(
        "payment_terms", "Invoiced Monthly \u2013 net 15 days on invoice"
    )
    p = doc.add_paragraph()
    p.add_run(f"[Payment Terms \u2013 {payment_terms}]").italic = True
    doc.add_paragraph()

    # ── Acceptance ────────────────────────────────────────────────
    h = doc.add_heading("Acceptance", level=1)
    _color_heading(h, 0x1B, 0x2B, 0x5E)
    doc.add_paragraph(ACCEPTANCE_BODY)
    doc.add_paragraph(CANCELLATION_CLAUSE)
    doc.add_paragraph(TERMINATION_CLAUSE)
    doc.add_paragraph()

    # ── Signature table ───────────────────────────────────────────
    sig = doc.add_table(rows=6, cols=3)
    sig.style = "Table Grid"

    # Header row
    sig.rows[0].cells[0].text = "CLIENT"
    sig.rows[0].cells[2].text = "TRACKABLEMED"
    _bold_cell(sig.rows[0].cells[0])
    _bold_cell(sig.rows[0].cells[2])
    _set_cell_width(sig.rows[0].cells[1], 0.3)

    sig_data = [
        ("Company", "", "TrackableMed"),
        ("Full Name", "", "Zed Williamson"),
        ("Title", "", "Founder/CEO"),
        ("Signature", "", ""),
        ("Date", "", ""),
    ]
    for i, (label, client_val, tm_val) in enumerate(sig_data):
        row = sig.rows[i + 1]
        row.cells[0].text = f"{label}:  {client_val}"
        row.cells[2].text = f"{label}:  {tm_val}"
        _set_cell_width(row.cells[1], 0.3)

    doc.add_paragraph()

    # ── Closing ───────────────────────────────────────────────────
    closing = doc.add_paragraph("I believe in the process!")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in closing.runs:
        run.bold = True
        run.italic = True

    # ── Return bytes ──────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PM Brief builder ─────────────────────────────────────────────

def build_pm_brief(fields: dict, content: dict) -> str:
    """
    Generate a plain-text PM brief for Camila to populate Wrike.
    Structured so she can copy-paste directly into task fields.
    """
    client = fields.get("client_name", "[Client Name]")
    contact = fields.get("client_contact", "")
    rep = fields.get("rep_name", "Erik Trattler")
    service = SERVICE_TYPES.get(fields.get("service_type", "seo_web"), "")
    today = date.today().strftime("%B %d, %Y")
    job_name = content.get("job_name", "")

    pricing_lines = "\n".join(
        f"  \u2022 {item['item']}: {item['cost']}"
        for item in fields.get("pricing_items", [])
    ) or "  \u2022 TBD \u2014 confirm with Erik"

    scope_bullets = "\n".join(
        f"  {i + 1}. {s['heading']}"
        for i, s in enumerate(content.get("scope_items", []))
    ) or "  \u2022 See attached SOW"

    timeline_section = ""
    if content.get("timeline_items"):
        lines = "\n".join(
            f"  \u2022 {t['activity']}: {t['start']} \u2013 {t['end']}"
            for t in content["timeline_items"]
        )
        timeline_section = f"\nTIMELINE\n{lines}\n"

    notes = fields.get("notes") or "None"
    goals = fields.get("client_goals") or "See SOW summary"

    return f"""PM BRIEF \u2014 {client}
Generated: {today}
\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501
ENGAGEMENT
  Client:         {client}
  Contact:        {contact}
  SOW Title:      {job_name}
  Service Type:   {service}
  Rep:            {rep}
  Goals:          {goals}

PRICING
{pricing_lines}

SCOPE (Wrike task breakdown)
{scope_bullets}
{timeline_section}
NOTES FOR KICKOFF
  {notes}

\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501\u2501
Schedule kickoff with Erik: https://meetings.hubspot.com/etrattler/meet-with-erik
"""
